import asyncio
import json
import os
import uuid
from contextlib import asynccontextmanager
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(override=True)

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from sse_starlette.sse import EventSourceResponse

from models.schemas import RunAgentsRequest, UploadPdfResponse, HealthResponse
from rag import store as rag_store
from rag.loader import load_pdf_chunks
from sessions import manager as session_manager
from orchestrator import run_pipeline


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Load startup PDF into ChromaDB
    pdf_path = os.getenv("PDF_PATH", "../Advanced Prompt Engineering [PROMO].pdf")
    if Path(pdf_path).exists():
        try:
            chunks = load_pdf_chunks(pdf_path)
            await rag_store.add_chunks(chunks, source="course_advanced_prompt_engineering")
            print(f"[startup] Loaded {len(chunks)} chunks from {pdf_path}")
        except Exception as e:
            print(f"[startup] PDF load warning: {e}")
    else:
        print(f"[startup] PDF not found at {pdf_path} — skipping RAG seed")

    # Start session cleanup background task
    asyncio.create_task(session_manager.cleanup_loop())
    yield


app = FastAPI(
    title="Supervisor Pattern Multi-Agent API",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        os.getenv("FRONTEND_URL", "http://localhost:3000"),
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── GET /health ────────────────────────────────────────────────────

@app.get("/health", response_model=HealthResponse)
async def health():
    return HealthResponse(status="ok", rag_chunks=rag_store.chunk_count())


# ── POST /upload-pdf ───────────────────────────────────────────────

@app.post("/upload-pdf", response_model=UploadPdfResponse)
async def upload_pdf(file: UploadFile = File(...)):
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted")

    import tempfile
    tmp_path = Path(tempfile.gettempdir()) / f"{uuid.uuid4()}_{file.filename}"
    try:
        content = await file.read()
        tmp_path.write_bytes(content)
        chunks = load_pdf_chunks(str(tmp_path))
        added = await rag_store.add_chunks(chunks, source=file.filename)
        return UploadPdfResponse(
            status="success",
            filename=file.filename,
            chunks=added,
            message=f"Processed {added} chunks from {file.filename}",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        tmp_path.unlink(missing_ok=True)


# ── POST /run-agents (SSE) ─────────────────────────────────────────

@app.post("/run-agents")
async def run_agents(request: RunAgentsRequest):
    session_id = request.session_id or str(uuid.uuid4())

    async def event_stream():
        # Send session ID first
        yield {"data": json.dumps({"type": "session_id", "session_id": session_id})}
        try:
            async for raw_event in run_pipeline(request, session_id):
                yield {"data": raw_event}
        except Exception as e:
            yield {"data": json.dumps({"type": "error", "message": str(e)})}

    return EventSourceResponse(event_stream())


# ── GET /download-report/{session_id}/{format} ─────────────────────

@app.get("/download-report/{session_id}/{fmt}")
async def download_report(session_id: str, fmt: str):
    state = session_manager.get(session_id)
    if not state or not state.final_report:
        raise HTTPException(status_code=404, detail="Report not found for this session")

    report = state.final_report

    if fmt == "md":
        return Response(
            content=report.encode("utf-8"),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="report_{session_id[:8]}.md"'},
        )
    elif fmt == "pdf":
        pdf_bytes = _report_to_pdf(report)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="report_{session_id[:8]}.pdf"'},
        )
    elif fmt == "docx":
        docx_bytes = _report_to_docx(report)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="report_{session_id[:8]}.docx"'},
        )
    else:
        raise HTTPException(status_code=400, detail="Format must be md, pdf, or docx")


# ── GET /session/{session_id}/state ────────────────────────────────

@app.get("/session/{session_id}/state")
async def get_session_state(session_id: str):
    state = session_manager.get(session_id)
    if not state:
        raise HTTPException(status_code=404, detail="Session not found")
    return state.to_dict()


# ── Export helpers ─────────────────────────────────────────────────

def _report_to_pdf(report: str) -> bytes:
    import textwrap
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Encode to latin-1 safe (FPDF core fonts limitation)
    safe = report.encode("latin-1", errors="replace").decode("latin-1")

    for line in safe.split("\n"):
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue

        # Headings
        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            text = stripped[4:]
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            text = stripped[3:]
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 15)
            text = stripped[2:]
        else:
            pdf.set_font("Helvetica", size=10)
            text = stripped

        # Wrap long tokens (code, URLs) at 95 chars so they always fit
        for wrapped_line in textwrap.wrap(text, width=95) or [text]:
            pdf.multi_cell(0, 5, wrapped_line, new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())


def _report_to_docx(report: str) -> bytes:
    import io
    from docx import Document
    doc = Document()
    doc.add_heading("Engineering Delivery Report", 0)
    for line in report.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
